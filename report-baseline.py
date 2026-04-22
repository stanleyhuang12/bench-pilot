"""
export_baseline.py
==================
Walks every benchmark sub-directory under results_root, finds test.json
and all conversation JSON files, then writes one baseline.docx per benchmark.

Usage
-----
    python export_baseline.py                          # auto-detects results/ or results-p/
    python export_baseline.py --results-root results-p
    python export_baseline.py --results-root results-p --benchmark abc-bench
    python export_baseline.py --out-name report.docx
"""

from __future__ import annotations

import argparse
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colours ───────────────────────────────────────────────────────────────────
class C:
    BLUE_DARK  = RGBColor(0x1F, 0x54, 0x96)
    BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)
    AMBER      = RGBColor(0xD4, 0x7A, 0x00)
    GREY       = RGBColor(0x66, 0x66, 0x66)
    WHITE      = "FFFFFF"
    BLUE_LIGHT = "EBF3FB"
    AMBER_BG   = "FFF3CD"
    HEADER_BG  = "2E75B6"


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _hr(doc: Document, color: str = "2E75B6", size: int = 6) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def _page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break()


# ── Paragraph helpers ─────────────────────────────────────────────────────────
def _heading(doc: Document, text: str, level: int = 1) -> None:
    colors_ = {1: C.BLUE_DARK, 2: C.BLUE_MID, 3: C.GREY}
    sizes_  = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes_.get(level, 12))
    run.font.color.rgb = colors_.get(level, C.GREY)


def _body(doc: Document, text: str, italic: bool = False,
          color: RGBColor | None = None, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = C.GREY
    r2 = p.add_run(value)
    r2.font.size = Pt(10)


# ── Table helpers ─────────────────────────────────────────────────────────────
def _make_table(doc: Document, col_widths_inches: list[float]):
    col_widths_dxa = [int(w * 1440) for w in col_widths_inches]
    table = doc.add_table(rows=0, cols=len(col_widths_dxa))
    table.style = "Table Grid"
    return table, col_widths_dxa


def _header_row(table, col_widths_dxa: list[int], labels: list[str]) -> None:
    row = table.add_row()
    for cell, label, w in zip(row.cells, labels, col_widths_dxa):
        _set_cell_bg(cell, C.HEADER_BG)
        _set_cell_margins(cell)
        cell.width = Twips(w)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _data_row(table, col_widths_dxa: list[int], values: list[str],
              alt: bool = False, bg_override: str | None = None,
              color_override: RGBColor | None = None) -> None:
    row = table.add_row()
    bg  = bg_override or (C.BLUE_LIGHT if alt else C.WHITE)
    for i, (cell, val, w) in enumerate(zip(row.cells, values, col_widths_dxa)):
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell)
        cell.width = Twips(w)
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9)
        if color_override and i == 0:
            run.font.color.rgb = color_override


# ── Truncate ──────────────────────────────────────────────────────────────────
def _truncate(text: str, first_n: int = 150, last_n: int = 150) -> str:
    words = text.split()
    if len(words) <= first_n + last_n:
        return text
    omitted = len(words) - first_n - last_n
    return (
        " ".join(words[:first_n])
        + f"\n\n[… {omitted} words omitted …]\n\n"
        + " ".join(words[-last_n:])
    )


# ── File discovery ────────────────────────────────────────────────────────────
def _find_json_files(directory: str) -> list[str]:
    """Recursively find all .json files under directory."""
    results = []
    for root, _, files in os.walk(directory):
        for fname in sorted(files):
            if fname.endswith(".json"):
                results.append(os.path.join(root, fname))
    return sorted(results)


def _auto_detect_root() -> str:
    for candidate in ("results", "results-p"):
        if os.path.isdir(candidate):
            return candidate
    return "results"


# ── Cover page ────────────────────────────────────────────────────────────────
def _write_cover(doc: Document, bench_name: str, test_data: dict,
                 n_convs: int) -> None:
    doc.add_paragraph()
    _hr(doc, "2E75B6", 12)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(4)
    run = title_p.add_run(
        test_data.get("benchmark_name", bench_name.replace("-", " ").title())
    )
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = C.BLUE_DARK

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run("Baseline Conversation Export")
    sub_r.font.size = Pt(13)
    sub_r.font.color.rgb = C.GREY
    sub_r.italic = True

    _hr(doc)
    doc.add_paragraph()

    _label_value(doc, "Benchmark",     bench_name)
    _label_value(doc, "Scenarios",     str(len(test_data.get("scenarios", []))))
    _label_value(doc, "Conversations", str(n_convs))
    _label_value(doc, "Generated",     datetime.now().strftime("%B %d, %Y"))

    desc = test_data.get("description", "")
    if desc:
        doc.add_paragraph()
        _body(doc, desc, italic=True, color=C.GREY)

    doc.add_paragraph()


# ── Scenario section ──────────────────────────────────────────────────────────
def _write_scenario(doc: Document, meta: dict, conv: dict) -> None:
    sid    = meta.get("id", "unknown")
    title  = meta.get("title", sid)
    lms    = meta.get("landmarks", [])
    has_lm = bool(lms)
    lm_map = {lm["turn"]: lm["instruction"] for lm in lms}

    # Demographic string
    demo = meta.get("demographic")
    if isinstance(demo, dict):
        demo_str = "  |  ".join(f"{k.capitalize()}: {v}" for k, v in demo.items())
    elif demo:
        demo_str = str(demo)
    else:
        demo_str = "—"

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(4)
    run = title_p.add_run(f"{'★ ' if has_lm else ''}{title}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C.BLUE_DARK

    _hr(doc, "2E75B6", 8)

    _label_value(doc, "Scenario ID",      sid)
    _label_value(doc, "Base Scenario ID", meta.get("base_scenario_id", "—"))
    _label_value(doc, "Demographic",      demo_str)
    doc.add_paragraph()

    _heading(doc, "Description", 2)
    _body(doc, meta.get("description", "—"))

    _heading(doc, "User Persona", 2)
    _body(doc, meta.get("user_persona", "—"))

    _heading(doc, "User Goal", 2)
    _body(doc, meta.get("user_goal", "—"))

    # Landmarks table
    if has_lm:
        _heading(doc, "Landmark Instructions  ★", 2)
        table, widths = _make_table(doc, [0.55, 8.81])
        _header_row(table, widths, ["Turn", "Instruction"])
        for i, lm in enumerate(lms):
            _data_row(
                table, widths,
                [str(lm.get("turn", "?")), lm.get("instruction", "")],
                alt=i % 2 == 1,
                bg_override=C.AMBER_BG if i % 2 == 0 else "FFFDE7",
                color_override=C.AMBER,
            )
        doc.add_paragraph()

    # Transcript
    _heading(doc, "Conversation Transcript", 2)
    _hr(doc)

    samples  = conv.get("samples")
    turns    = (samples[0] if samples else None) or conv.get("turns", [])
    n_samples = len(samples) if samples else 1

    if n_samples > 1:
        _body(doc, f"(Showing sample 1 of {n_samples})", italic=True, color=C.GREY)

    if not turns:
        _body(doc, "(No conversation turns found)", italic=True, color=C.GREY)
    else:
        for t_idx, turn in enumerate(turns):
            role     = turn.get("role", "")
            content  = turn.get("content", "")
            turn_no  = t_idx // 2 + 1
            is_user  = role == "user"
            is_lm    = is_user and turn_no in lm_map
            role_str = "USER" if is_user else "ASSISTANT"
            prefix   = "★ " if is_lm else ""

            # Role label
            label_p = doc.add_paragraph()
            label_p.paragraph_format.space_before = Pt(10 if is_lm else 6)
            label_p.paragraph_format.space_after  = Pt(1)
            label_r = label_p.add_run(f"{prefix}Turn {turn_no}  —  {role_str}")
            label_r.bold = True
            label_r.font.size = Pt(9)
            label_r.font.color.rgb = (
                C.AMBER if is_lm else
                C.BLUE_MID if not is_user else
                C.GREY
            )

            # Landmark callout
            if is_lm:
                inst_p = doc.add_paragraph()
                inst_p.paragraph_format.left_indent  = Inches(0.2)
                inst_p.paragraph_format.space_before = Pt(0)
                inst_p.paragraph_format.space_after  = Pt(2)
                inst_r = inst_p.add_run(f"Landmark: {lm_map[turn_no]}")
                inst_r.font.size = Pt(8)
                inst_r.italic = True
                inst_r.font.color.rgb = C.AMBER

            display = content if is_user else _truncate(content)
            body_p = doc.add_paragraph()
            body_p.paragraph_format.left_indent  = Inches(0.2)
            body_p.paragraph_format.space_before = Pt(0)
            body_p.paragraph_format.space_after  = Pt(10 if is_lm else 4)
            run = body_p.add_run(display)
            run.font.size = Pt(9)
            if is_lm:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

def _write_metrics(doc: Document, test_data: dict) -> None:
    metrics = test_data.get("metrics", [])
    if not metrics:
        return

    _heading(doc, "Evaluation Metrics", 1)
    _hr(doc)

    table, widths = _make_table(doc, [2.0, 2.5, 1.0, 1.5, 3.0])
    _header_row(table, widths, ["ID", "Name", "Type", "Applies To", "Description"])

    for i, m in enumerate(metrics):
        desc = (m.get("description", "") or "").split("\n")[0]  # keep concise
        _data_row(
            table,
            widths,
            [
                m.get("id", ""),
                m.get("name", ""),
                m.get("type", ""),
                m.get("applies_to", ""),
                desc,
            ],
            alt=i % 2 == 1,
        )

    doc.add_paragraph()
    
# ── Benchmark builder ─────────────────────────────────────────────────────────
def build_benchmark(bench_dir: str, bench_name: str, out_name: str, max_convs: int) -> None:
    test_path = os.path.join(bench_dir, "test.json")
    if not os.path.exists(test_path):
        print(f"  ⚠  No test.json in {bench_dir} — skipping")
        return

    with open(test_path) as f:
        test_data = json.load(f)

    scenario_lookup = {s["id"]: s for s in test_data.get("scenarios", [])}

    # Find conversation JSONs recursively under conversations/ (or bench_dir)
    conv_base = os.path.join(bench_dir, "conversations")
    search_root = conv_base if os.path.isdir(conv_base) else bench_dir
    conv_files = [
        p for p in _find_json_files(search_root)
        if not p.endswith("test.json") and not p.endswith("results.json")
    ]

    if not conv_files:
        print(f"  ⚠  No conversation files found under {search_root}")
        return

    print(f"  Found {len(conv_files)} conversation file(s)")

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Twips(12240)
    section.page_height = Twips(15840)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(1))
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _write_cover(doc, bench_name, test_data, len(conv_files))
    _write_metrics(doc, test_data)
    _page_break(doc)

    first = True
    count = 0

    for fpath in conv_files:
        if count >= max_convs:
            break

        with open(fpath) as f:
            data = json.load(f)

        convs = data if isinstance(data, list) else [data]

        for conv in convs:
            if count >= max_convs:
                break

            if not isinstance(conv, dict):
                continue

            sid = conv.get("scenario_id") or os.path.basename(fpath).replace(".json", "")
            meta = scenario_lookup.get(sid) or conv.get("scenario")

            if not meta:
                continue

            if not first:
                _page_break(doc)
            first = False

            _write_scenario(doc, meta, conv)
            print(f"    ✓ {sid}")

            count += 1

            _write_scenario(doc, meta, conv)
            print(f"    ✓ {sid}")

    out_path = os.path.join(bench_dir, out_name)
    doc.save(out_path)
    print(f"  ✓ Saved → {out_path}")


# ── Main ──────────────────────────────────────────────────────────────────────
def main() -> None:
    parser = argparse.ArgumentParser(description="Export baseline conversations to .docx")
    parser.add_argument("--results-root", default=None,
                        help="Root directory containing benchmark folders (default: auto-detect results/ or results-p/)")
    parser.add_argument("--benchmark", default=None,
                        help="Run only this benchmark subdirectory")
    parser.add_argument("--out-name", default="baseline.docx",
                        help="Output filename written into each benchmark folder (default: baseline.docx)")
    parser.add_argument("--max-convs", type=int, default=3,
                    help="Maximum number of conversations to include (default: 3)")
    args = parser.parse_args()

    results_root = args.results_root or _auto_detect_root()

    if not os.path.isdir(results_root):
        raise SystemExit(f"Results root not found: {results_root}")

    benchmarks = [
        name for name in sorted(os.listdir(results_root))
        if os.path.isdir(os.path.join(results_root, name))
        and (args.benchmark is None or name == args.benchmark)
    ]

    if not benchmarks:
        raise SystemExit(f"No benchmark directories found in {results_root}")

    print(f"\nResults root : {results_root}")
    print(f"Output file  : {args.out_name}")
    print(f"Benchmarks   : {', '.join(benchmarks)}\n")

    for bench in benchmarks:
        bench_dir = os.path.join(results_root, bench)
        print(f"\n── {bench} ──")
        build_benchmark(bench_dir, bench, args.out_name, args.max_convs)

    print("\nDone.\n")


if __name__ == "__main__":
    main()
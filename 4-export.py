"""
4-export.py — Phase 4: Compile all benchmark data into a results.docx report.

Usage:
    python 4-export.py
    python 4-export.py --config path/to/config.json
    python 4-export.py --results-root results --benchmark emotional-dependency
    python 4-export.py --output report.docx
"""

from __future__ import annotations

import argparse
import json
import os
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import load_config, get_model_name


class C:
    ACCENT      = RGBColor(0x2E, 0x75, 0xB6)
    ACCENT_DARK = RGBColor(0x1F, 0x54, 0x96)
    PASS        = RGBColor(0x1E, 0x84, 0x49)
    FAIL        = RGBColor(0xC0, 0x39, 0x2B)
    WARN        = RGBColor(0xD4, 0x7A, 0x00)   # amber — landmark highlight
    HEADER_BG   = "2E75B6"
    LANDMARK_BG = "FFF3CD"                      # warm yellow — landmark turn
    ROW_ALT     = "EBF3FB"
    WHITE       = "FFFFFF"
    MID_GREY    = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _add_horizontal_rule(doc: Document, color: str = "2E75B6", size: int = 6) -> None:
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    colors = {1: C.ACCENT_DARK, 2: C.ACCENT, 3: C.MID_GREY}
    sizes  = {1: 20, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(sizes.get(level, 12))
    run.font.color.rgb = colors.get(level, C.ACCENT)


def _body(
    doc: Document,
    text: str,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.italic     = italic
    if color:
        run.font.color.rgb = color


def _label_value(doc: Document, label: str, value: str) -> None:
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{label}: ")
    r1.bold            = True
    r1.font.size       = Pt(10)
    r1.font.color.rgb  = C.MID_GREY
    r2 = p.add_run(value)
    r2.font.size       = Pt(10)


def _make_table(doc: Document, col_widths_inches: list[float]):
    col_widths_dxa = [int(w * 1440) for w in col_widths_inches]
    table = doc.add_table(rows=0, cols=len(col_widths_dxa))
    table.style = "Table Grid"
    return table, col_widths_dxa


def _header_row(table, col_widths_dxa: list[int], labels: list[str]) -> None:
    row = table.add_row()
    for cell, label, w in zip(row.cells, labels, col_widths_dxa):
        _set_cell_bg(cell, C.HEADER_BG)
        _set_cell_margins(cell)
        cell.width = Twips(w)
        run = cell.paragraphs[0].add_run(label)
        run.bold            = True
        run.font.size       = Pt(9)
        run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)


def _data_row(
    table,
    col_widths_dxa: list[int],
    values: list[str],
    alt: bool = False,
    bold_first: bool = False,
    verdict_col: int | None = None,
    bg_override: str | None = None,   # hex, e.g. C.LANDMARK_BG
) -> None:
    row = table.add_row()
    bg  = bg_override or (C.ROW_ALT if alt else C.WHITE)
    for i, (cell, val, w) in enumerate(zip(row.cells, values, col_widths_dxa)):
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell)
        cell.width = Twips(w)
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9)
        if i == 0 and bold_first:
            run.bold = True
        if verdict_col is not None and i == verdict_col:
            low = str(val).lower()
            run.font.color.rgb = C.PASS if low == "yes" else C.FAIL
            run.bold = True


# ---------------------------------------------------------------------------
# Data loading helpers
# ---------------------------------------------------------------------------

def _resolve_benchmarks(
    results_root: str,
    test_path: str,
    benchmark: str | None,
) -> list[str]:
    entries = os.listdir(results_root)
    if benchmark:
        if benchmark not in entries:
            raise FileNotFoundError(f"Benchmark '{benchmark}' not found in {results_root}")
        return [os.path.join(results_root, benchmark, test_path)]
    return [os.path.join(results_root, b, test_path) for b in entries]


def _load_benchmark(bench: str, conv_dir_name: str, results_filename: str) -> dict:
    bench_dir  = os.path.dirname(bench)
    bench_name = os.path.basename(bench_dir)
    conv_path  = os.path.join(bench_dir, conv_dir_name)
    results_path = os.path.join(bench_dir, results_filename)

    for path, label in [
        (bench,        "generate.py"),
        (conv_path,    "simulate.py"),
        (results_path, "evaluate.py"),
    ]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"{path} — run {label} first.")

    with open(bench) as f:
        test_data = json.load(f)

    conversations: dict[str, dict] = {}
    for fname in sorted(os.listdir(conv_path)):
        if fname.endswith(".json"):
            with open(os.path.join(conv_path, fname)) as f:
                conv = json.load(f)
                sid  = conv.get("scenario_id") or fname.replace(".json", "")
                conversations[sid] = conv

    with open(results_path) as f:
        results = json.load(f)

    return {
        "name":          bench_name,
        "test_data":     test_data,
        "conversations": conversations,
        "results":       results,
    }


def _load_goal_content(config: dict) -> str:
    goal_path = config.get("paths", {}).get("goal_prompt", "")
    if goal_path and os.path.exists(goal_path):
        with open(goal_path) as f:
            return f.read().strip()
    return "(goal file not found)"


def _num_iterations(bench: dict) -> int:
    for d in bench["results"].get("details", []):
        r = d.get("results", [])
        if isinstance(r, list) and r:
            return len(r)
    return 1


# ---------------------------------------------------------------------------
# Detail record accessors
# ---------------------------------------------------------------------------

def _detail_result(detail: dict, sample: int = 0) -> str:
    results = detail.get("results", [])
    if isinstance(results, list) and sample < len(results):
        return str(results[sample]).lower()
    if isinstance(results, str):
        return results.lower()
    return "—"


def _detail_justification(detail: dict, sample: int = 0) -> str:
    justs = detail.get("justifications", [])
    if isinstance(justs, list) and sample < len(justs):
        return justs[sample]
    if isinstance(justs, str):
        return justs
    return detail.get("justification", "")


# ---------------------------------------------------------------------------
# Landmark helpers
# ---------------------------------------------------------------------------

def _landmark_turn_map(scenario: dict) -> dict[int, str]:
    """Return {1-indexed turn number: instruction} for all landmarks."""
    return {lm["turn"]: lm["instruction"] for lm in scenario.get("landmarks", [])}


def _has_landmarks(scenario: dict) -> bool:
    return bool(scenario.get("landmarks"))


def _truncate_message(text: str, first_n: int = 150, last_n: int = 150) -> str:
    """
    Keep at most first_n + last_n words of a message.
    If the message is short enough, return it unchanged.
    When truncated, a centred ellipsis line shows how many words were cut.
    """
    words = text.split()
    total = len(words)
    if total <= first_n + last_n:
        return text
    omitted = total - first_n - last_n
    head = " ".join(words[:first_n])
    tail = " ".join(words[-last_n:])
    return f"{head}\n\n[… {omitted} words omitted …]\n\n{tail}"


def _select_variants(
    scenarios: list[dict],
    by_scenario: dict,
    variants_per_base: int = 2,
) -> list[dict]:
    """
    Group all scenario variants by their base_scenario_id (falling back to
    the scenario's own id for base / non-expanded scenarios), then pick the
    top ``variants_per_base`` within each group ranked by harm_rate descending.

    If no harm data is available for a variant it is treated as harm_rate = 0
    and floats to the bottom of each group (so high-harm cases are always shown).

    Returns a flat list preserving the sorted order within each group, with
    groups themselves in base_scenario_id alphabetical order.
    """
    from collections import defaultdict

    groups: dict[str, list[dict]] = defaultdict(list)
    for sc in scenarios:
        base_id = sc.get("base_scenario_id") or sc["id"]
        groups[base_id].append(sc)

    selected: list[dict] = []
    for base_id in sorted(groups):
        variants = groups[base_id]
        ranked = sorted(
            variants,
            key=lambda sc: by_scenario.get(sc["id"], {}).get("harm_rate", 0.0),
            reverse=True,
        )
        selected.extend(ranked[:variants_per_base])

    return selected


def _build_base_scenario_lookup(scenarios: list[dict]) -> dict[str, str]:
    """
    Map base_scenario_id → title using the first variant that references it.
    Also handles plain (non-expanded) scenarios where id == base_scenario_id.
    """
    lookup: dict[str, str] = {}
    for sc in scenarios:
        base_id = sc.get("base_scenario_id") or sc["id"]
        if base_id not in lookup:
            lookup[base_id] = sc.get("title", base_id)
    return lookup


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------

def _write_cover(doc, bench, config, user_model, target_model) -> None:
    test_data = bench["test_data"]
    n_iter    = _num_iterations(bench)
    s         = bench["results"].get("summary", {})

    doc.add_paragraph()
    _add_horizontal_rule(doc, color="2E75B6", size=12)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(4)
    run = title_p.add_run(
        test_data.get("benchmark_name", bench["name"].replace("-", " ").title())
    )
    run.bold            = True
    run.font.size       = Pt(28)
    run.font.color.rgb  = C.ACCENT_DARK

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run("AI Conversation Evaluation Report")
    sub_r.font.size    = Pt(13)
    sub_r.font.color.rgb = C.MID_GREY
    sub_r.italic       = True

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    _label_value(doc, "Description",              test_data.get("description", "—"))
    _label_value(doc, "User model",               user_model)
    _label_value(doc, "Target model",             target_model)
    _label_value(doc, "Generated",                datetime.now().strftime("%B %d, %Y"))
    _label_value(doc, "Scenarios",                str(len(bench["conversations"])))
    _label_value(doc, "Metrics",                  str(len(test_data.get("metrics", []))))
    _label_value(doc, "Iterations per scenario",  str(n_iter))
    _label_value(doc, "Overall yes rate",         f"{s.get('yes_rate', 0):.1%}")
    _label_value(doc, "Overall harm rate",        f"{s.get('harm_rate', 0):.1%}")

    doc.add_paragraph()


def _write_model_and_prompts(doc, config, goal_content) -> None:
    _heading(doc, "Model Configuration & Prompts", level=1)
    _add_horizontal_rule(doc)

    _heading(doc, "Models", level=2)
    models = config.get("models", {})
    table, widths = _make_table(doc, [1.5, 3.5, 4.0])
    _header_row(table, widths, ["Role", "Provider / Model", "Notes"])
    for i, (role, label) in enumerate(
        {"generator": "Generator", "user": "User Simulator", "target": "Target"}.items()
    ):
        m = models.get(role, {})
        _data_row(
            table, widths,
            [label, f"{m.get('provider', '—')} / {m.get('model', '—')}", m.get("notes", "")],
            alt=i % 2 == 1,
        )
    doc.add_paragraph()

    _heading(doc, "Generation Parameters", level=2)
    gen   = config.get("generation", {})
    paths = config.get("paths", {})
    table2, widths2 = _make_table(doc, [3.0, 6.36])
    _header_row(table2, widths2, ["Parameter", "Value"])
    for i, (k, v) in enumerate([
        ("Scenarios per batch",    str(gen.get("num_scenarios", "—"))),
        ("Turns per conversation", str(gen.get("turns_per_conversation", "—"))),
        ("Goal prompt file",       paths.get("goal_prompt", "—")),
        ("Test file",              paths.get("test_file", "—")),
        ("Conversations dir",      paths.get("conversations_dir", "—")),
        ("Results file",           paths.get("results_file", "—")),
    ]):
        _data_row(table2, widths2, [k, v], alt=i % 2 == 1)
    doc.add_paragraph()

    _heading(doc, "Goal Prompt (goal.json / goal.md)", level=2)
    _body(doc, "Raw content of the goal file used for scenario and metric generation.",
          italic=True, color=C.MID_GREY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Inches(0.2)
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(6)
    run2 = p2.add_run(goal_content)
    run2.font.name       = "Courier New"
    run2.font.size       = Pt(8)
    run2.font.color.rgb  = C.MID_GREY
    doc.add_paragraph()


def _write_metrics_overview(doc, bench) -> None:
    metrics = bench["test_data"].get("metrics", [])
    _heading(doc, "Evaluated Metrics", level=1)
    _add_horizontal_rule(doc)

    table, widths = _make_table(doc, [1.0, 1.6, 3.8, 1.2, 1.76])
    _header_row(table, widths, ["ID", "Name", "Description", "Type", "Applies To"])
    for i, m in enumerate(metrics):
        applies = m.get("applies_to", "all")
        applies_str = "All scenarios" if applies == "all" else ", ".join(applies)
        _data_row(
            table, widths,
            [
                m.get("id", ""),
                m.get("name", ""),
                m.get("description", ""),
                m.get("type", ""),
                applies_str,
            ],
            alt=i % 2 == 1,
        )
    doc.add_paragraph()


def _write_aggregate(doc, bench) -> None:
    """
    Render summary statistics and per-metric / per-scenario roll-ups.

    Uses pre-computed fields from results.json:
      summary.yes_rate, summary.harm_rate, summary.total_yes, summary.total_valid
      by_metric[mid].yes_rate / harm_rate / yes / valid / percent_agreement
      by_scenario[sid].yes_rate / harm_rate
    """
    results = bench["results"]
    s       = results.get("summary", {})

    _heading(doc, "Aggregate Results", level=1)
    _add_horizontal_rule(doc)

    # ── Headline rates ─────────────────────────────────────────────────────
    yes_rate  = s.get("yes_rate",  0.0)
    harm_rate = s.get("harm_rate", 0.0)
    total_yes   = s.get("total_yes",   0)
    total_valid = s.get("total_valid", 0)
    total_harm  = s.get("total_harm",  0)

    # Progress bar (yes rate)
    filled = round(yes_rate * 20)
    bar_p  = doc.add_paragraph()
    bar_p.paragraph_format.space_after = Pt(2)
    bar_r  = bar_p.add_run("█" * filled)
    bar_r.font.color.rgb = C.PASS
    bar_r.font.size      = Pt(14)
    empty_r = bar_p.add_run("█" * (20 - filled))
    empty_r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    empty_r.font.size      = Pt(14)
    label_r = bar_p.add_run(f"  {yes_rate:.1%}  ({total_yes}/{total_valid} yes)")
    label_r.font.size      = Pt(11)
    label_r.bold           = True
    label_r.font.color.rgb = C.PASS if yes_rate >= 0.5 else C.FAIL

    # Harm rate line
    harm_p = doc.add_paragraph()
    harm_p.paragraph_format.space_before = Pt(2)
    harm_p.paragraph_format.space_after  = Pt(8)
    h_label = harm_p.add_run("Harm rate: ")
    h_label.bold = True
    h_label.font.size = Pt(10)
    h_val = harm_p.add_run(f"{harm_rate:.1%}  ({total_harm}/{total_valid} harmful responses)")
    h_val.font.size      = Pt(10)
    h_val.font.color.rgb = C.FAIL if harm_rate > 0.1 else C.PASS

    # ── By-metric table ────────────────────────────────────────────────────
    _heading(doc, "By Metric", level=2)
    by_metric     = results.get("by_metric", {})
    metric_lookup = {m["id"]: m["name"] for m in bench["test_data"].get("metrics", [])}
    table, widths = _make_table(doc, [1.2, 2.4, 1.1, 1.1, 0.9, 1.1, 1.56])
    _header_row(table, widths, ["Metric ID", "Name", "Yes Rate", "Harm Rate", "Yes", "Valid", "Agreement %"])
    for i, (mid, v) in enumerate(sorted(by_metric.items())):
        pa = v.get("percent_agreement")
        _data_row(table, widths, [
            mid,
            metric_lookup.get(mid, mid),
            f"{v.get('yes_rate',  0):.1%}",
            f"{v.get('harm_rate', 0):.1%}",
            str(v.get("yes",   0)),
            str(v.get("valid", 0)),
            f"{pa:.0%}" if pa is not None else "—",
        ], alt=i % 2 == 1, bold_first=True)
    doc.add_paragraph()

    # ── By-base-scenario table ─────────────────────────────────────────────
    # Rolls up all demographic variants into a single row per base scenario,
    # giving a clean topic-level view without per-variant noise.
    _heading(doc, "By Base Scenario", level=2)
    by_base = results.get("by_base_scenario", {})
    base_title_lookup = _build_base_scenario_lookup(bench["test_data"].get("scenarios", []))
    table2, widths2 = _make_table(doc, [1.3, 2.7, 1.1, 1.1, 0.9, 1.1, 1.16])
    _header_row(table2, widths2, ["Base ID", "Title", "Yes Rate", "Harm Rate", "Yes", "Valid", "Agreement %"])
    for i, (base_id, v) in enumerate(sorted(by_base.items())):
        pa = v.get("percent_agreement")
        _data_row(table2, widths2, [
            base_id,
            base_title_lookup.get(base_id, base_id),
            f"{v.get('yes_rate',  0):.1%}",
            f"{v.get('harm_rate', 0):.1%}",
            str(v.get("yes",   0)),
            str(v.get("valid", 0)),
            f"{pa:.0%}" if pa is not None else "—",
        ], alt=i % 2 == 1, bold_first=True)
    doc.add_paragraph()


def _write_landmark_transcript(
    doc: Document,
    scenario: dict,
    samples: list[list[dict]],
    n_iter: int,
) -> None:
    """
    Render the Sample 1 transcript, highlighting turns that correspond to
    a landmark instruction.  Landmark turns are shaded amber with the
    instruction shown inline.

    Only called for scenarios that have at least one landmark.
    """
    landmark_map = _landmark_turn_map(scenario)
    if not samples:
        return

    _heading(doc, f"Conversation Transcript (Sample 1 of {n_iter})", level=3)

    # Legend
    legend_p = doc.add_paragraph()
    legend_p.paragraph_format.space_before = Pt(2)
    legend_p.paragraph_format.space_after  = Pt(6)
    leg_run = legend_p.add_run("★  Shaded turns are landmark turns (special simulator instructions were active)")
    leg_run.font.size      = Pt(8)
    leg_run.italic         = True
    leg_run.font.color.rgb = C.WARN

    turns = samples[0]
    for t_idx, turn in enumerate(turns):
        role        = turn.get("role", "")
        content     = turn.get("content", "")
        turn_number = t_idx // 2 + 1   # 1-indexed conversation turn (user+assistant = 1 turn)
        is_landmark = turn_number in landmark_map and role == "user"

        # Role label
        label_p = doc.add_paragraph()
        label_p.paragraph_format.space_before = Pt(6 if not is_landmark else 10)
        label_p.paragraph_format.space_after  = Pt(1)
        role_str = "USER" if role == "user" else "ASSISTANT"
        label_r  = label_p.add_run(
            f"{'★ ' if is_landmark else ''}Turn {turn_number}  —  {role_str}"
        )
        label_r.bold       = True
        label_r.font.size  = Pt(9)
        label_r.font.color.rgb = (
            C.WARN if is_landmark else
            C.ACCENT if role == "assistant" else
            C.MID_GREY
        )

        # Landmark instruction callout (above the message)
        if is_landmark:
            inst_p = doc.add_paragraph()
            inst_p.paragraph_format.left_indent  = Inches(0.2)
            inst_p.paragraph_format.space_before = Pt(0)
            inst_p.paragraph_format.space_after  = Pt(2)
            inst_r = inst_p.add_run(f"Landmark instruction: {landmark_map[turn_number]}")
            inst_r.font.size       = Pt(8)
            inst_r.italic          = True
            inst_r.font.color.rgb  = C.WARN

        # Message body — assistant messages are truncated to avoid wall-of-text
        display_content = (
            _truncate_message(content)
            if role == "assistant"
            else content
        )
        body_p = doc.add_paragraph()
        body_p.paragraph_format.left_indent  = Inches(0.2)
        body_p.paragraph_format.space_before = Pt(0)
        body_p.paragraph_format.space_after  = Pt(4 if not is_landmark else 10)
        run = body_p.add_run(display_content)
        run.font.size = Pt(9)
        if is_landmark:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _write_scenarios(
    doc,
    bench,
    variants_per_base: int = 2,
) -> None:
    """
    Detailed per-variant section.

    Selection strategy
    ------------------
    Scenarios are grouped by ``base_scenario_id``.  Within each group the
    ``variants_per_base`` variants with the highest harm_rate are kept; all
    others are silently skipped.  This keeps the report to a manageable size
    even when there are many demographic combinations (e.g. 24 per base).

    Transcript policy
    -----------------
    - Transcripts are shown ONLY for scenarios that contain landmark turns.
    - Non-landmark scenarios show a one-line summary instead.
    - Assistant messages are always truncated to ≤150 words head + 150 tail.
    """
    results     = bench["results"]
    convs       = bench["conversations"]
    metrics     = bench["test_data"].get("metrics", [])
    all_details = results.get("details", [])
    n_iter      = _num_iterations(bench)
    by_scenario = results.get("by_scenario", {})

    # Select the top-harm variants per base scenario
    all_scenarios  = bench["test_data"].get("scenarios", [])
    selected       = _select_variants(all_scenarios, by_scenario, variants_per_base)
    selected_ids   = {sc["id"] for sc in selected}
    total_variants = len(all_scenarios)

    # Build (scenario_id, metric_id) → detail record index
    detail_map: dict[tuple, dict] = {}
    for d in all_details:
        if not isinstance(d, dict):
            continue
        sid_key = d.get("scenario_id")
        mid_key = d.get("metric_id") or d.get("metric")
        if sid_key and mid_key:
            detail_map[(sid_key, mid_key)] = d

    _heading(doc, "Scenario Results", level=1)
    _add_horizontal_rule(doc)
    _body(
        doc,
        f"Showing {len(selected)} of {total_variants} variants "
        f"({variants_per_base} per base scenario, ranked by harm rate descending).  "
        "Transcripts are included only for variants with landmark turns (★).  "
        "Assistant messages are truncated to the first and last 150 words.",
        italic=True, color=C.MID_GREY,
    )

    # Track which base scenario we last printed so we can add a group header
    last_base_id: str | None = None

    for scenario in selected:
        sid          = scenario["id"]
        base_id      = scenario.get("base_scenario_id") or sid
        conv         = convs.get(sid, {})
        samples      = conv.get("samples") or (
            [conv.get("turns", [])] if conv.get("turns") else []
        )
        has_lm       = _has_landmarks(scenario)
        variant_harm = by_scenario.get(sid, {}).get("harm_rate", 0.0)

        # ── Base-scenario group header (printed once per group) ────────────
        if base_id != last_base_id:
            _add_page_break(doc)
            _heading(doc, f"Base Scenario: {base_id}", level=1)
            _add_horizontal_rule(doc, size=4)
            last_base_id = base_id

        # ── Variant header ─────────────────────────────────────────────────
        title_str = f"{'★ ' if has_lm else ''}{sid}  —  {scenario['title']}"
        _heading(doc, title_str, level=2)

        _label_value(doc, "Description",  scenario.get("description", ""))
        _label_value(doc, "User goal",    scenario.get("user_goal", ""))
        if scenario.get("user_persona"):
            _label_value(doc, "User persona", scenario["user_persona"])
        if scenario.get("demographic"):
            demo = scenario["demographic"]
            demo_str = "  |  ".join(
                f"{k.capitalize()}: {v}" for k, v in demo.items()
            )
            _label_value(doc, "Demographic", demo_str)
        if has_lm:
            lm_strs = [
                f"Turn {lm['turn']}: {lm['instruction']}"
                for lm in scenario["landmarks"]
            ]
            _label_value(doc, "Landmarks", "  ·  ".join(lm_strs))
        _label_value(doc, "Harm rate (this variant)", f"{variant_harm:.1%}")
        _label_value(doc, "Iterations",               str(n_iter))
        doc.add_paragraph()

        # ── Agreement across iterations ────────────────────────────────────
        _heading(doc, "Agreement Across Iterations", level=3)
        table_a, widths_a = _make_table(doc, [1.3, 2.5, 1.3, 0.7, 0.7, 2.86])
        _header_row(table_a, widths_a, ["Metric ID", "Name", "Agreement %", "Yes", "No", "Consistency"])

        for i, metric in enumerate(metrics):
            mid    = metric["id"]
            detail = detail_map.get((sid, mid))

            if detail is None:
                row = table_a.add_row()
                bg  = C.ROW_ALT if i % 2 == 1 else C.WHITE
                for cell, val, w in zip(
                    row.cells,
                    [mid, metric.get("name", mid), "—", "—", "—", "no data"],
                    widths_a,
                ):
                    _set_cell_bg(cell, bg)
                    _set_cell_margins(cell)
                    cell.width = Twips(w)
                    cell.paragraphs[0].add_run(val).font.size = Pt(9)
                continue

            result_list = detail.get("results", [])
            if isinstance(result_list, str):
                result_list = [result_list]
            result_list = [str(r).lower() for r in result_list]

            n_yes   = sum(1 for r in result_list if r == "yes")
            n_no    = len(result_list) - n_yes
            total   = len(result_list)
            n_agree   = max(n_yes, n_no)
            agree_pct = n_agree / total if total else 0.0
            consistency = (
                "High"     if agree_pct >= 0.9 else
                "Moderate" if agree_pct >= 0.6 else
                "Low"
            )

            row = table_a.add_row()
            bg  = C.ROW_ALT if i % 2 == 1 else C.WHITE
            vals = [
                mid, metric.get("name", mid),
                f"{agree_pct:.0%}", str(n_yes), str(n_no), consistency,
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, widths_a)):
                _set_cell_bg(cell, bg)
                _set_cell_margins(cell)
                cell.width = Twips(w)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
                if j == 0:
                    run.bold = True
                if j == 2:
                    run.font.color.rgb = C.PASS if agree_pct >= 0.6 else C.FAIL
                    run.bold = True
                if j == 5:
                    run.font.color.rgb = (
                        C.PASS if consistency == "High"
                        else C.FAIL if consistency == "Low"
                        else C.MID_GREY
                    )
        doc.add_paragraph()

        # ── Metric scorecard (Sample 1) ────────────────────────────────────
        _heading(doc, f"Metric Scorecard (Sample 1 of {n_iter})", level=3)
        table_s, widths_s = _make_table(doc, [1.3, 2.2, 0.85, 5.01])
        _header_row(table_s, widths_s, ["Metric ID", "Name", "Result", "Justification (Sample 1)"])

        for i, metric in enumerate(metrics):
            mid    = metric["id"]
            detail = detail_map.get((sid, mid))
            if detail is None:
                continue
            verdict = _detail_result(detail, sample=0)
            just    = _detail_justification(detail, sample=0)
            _data_row(
                table_s, widths_s,
                [mid, metric.get("name", mid), verdict.upper(), just],
                alt=i % 2 == 1,
                verdict_col=2,
            )
        doc.add_paragraph()

        # ── Transcript (landmark variants only) ────────────────────────────
        if has_lm and samples:
            _write_landmark_transcript(doc, scenario, samples, n_iter)
        else:
            note_p = doc.add_paragraph()
            note_p.paragraph_format.space_before = Pt(2)
            note_p.paragraph_format.space_after  = Pt(6)
            note_r = note_p.add_run(
                "No landmark turns — transcript omitted.  "
                "Raw conversation files are in the conversations/ directory."
            )
            note_r.font.size      = Pt(9)
            note_r.italic         = True
            note_r.font.color.rgb = C.MID_GREY

        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def export(
    config_path: str = "config.json",
    results_root: str = "results",
    benchmark: str | None = None,
    output_path: str = "report.docx",
    variants_per_base: int = 2,
) -> None:
    config       = load_config(config_path)
    test_path    = config["paths"]["test_file"]
    conv_dir     = config["paths"]["conversations_dir"]
    results_file = config["paths"]["results_file"]
    user_model   = get_model_name(config, "user")
    target_model = get_model_name(config, "target")
    goal_content = _load_goal_content(config)

    bench_paths = _resolve_benchmarks(results_root, test_path, benchmark)

    for bench_path in bench_paths:
        bench = _load_benchmark(bench_path, conv_dir, results_file)
        out   = (
            os.path.join(results_root, benchmark, "report.docx")
            if len(bench_paths) > 1 or output_path == "report.docx"
            else output_path
        )

        print(f"\n  Building report for: {bench['name']}")

        doc     = Document()
        section = doc.sections[0]
        section.page_width  = Twips(12240)
        section.page_height = Twips(15840)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Inches(1))
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)

        _write_cover(doc, bench, config, user_model, target_model)
        _add_page_break(doc)
        _write_model_and_prompts(doc, config, goal_content)
        _add_page_break(doc)
        _write_metrics_overview(doc, bench)
        _add_page_break(doc)
        _write_aggregate(doc, bench)
        _write_scenarios(doc, bench, variants_per_base=variants_per_base)

        doc.save(out)
        print(f"  ✓ Saved → {out}")

    print("\nDone.\n")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Phase 4: Export results to .docx")
    parser.add_argument("--config",       default="config.json")
    parser.add_argument("--results-root", default="results")
    parser.add_argument("--benchmark",    type=str, required=False)
    parser.add_argument(
        "--output", type=str, default="report.docx",
        help="Output path (ignored when running all benchmarks)",
    )
    parser.add_argument(
        "--variants-per-base",
        type=int,
        default=2,
        metavar="N",
        help=(
            "Number of demographic variants to show per base scenario in the "
            "detailed section, ranked by harm rate descending (default: 2).  "
            "Pass 0 to show all variants."
        ),
    )
    args = parser.parse_args()

    export(
        config_path=args.config,
        results_root=args.results_root,
        benchmark=args.benchmark,
        output_path=args.output,
        variants_per_base=args.variants_per_base if args.variants_per_base > 0 else 10_000,
    )